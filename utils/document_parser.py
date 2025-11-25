"""
Document Parser Utility
Supports parsing PDF, DOCX, DOC files using MinerU and python-docx
"""

import os
import re
import tempfile
from typing import Optional, Dict, Any, List, Tuple
from pathlib import Path

from utils.logger import logger

FORMULA_HINT_CHARS = set("=±×÷∑∏√∞πσΔΩμθλκ^_{}<>≠≥≤")
TABLE_DELIMITERS = ("|", "\t")

try:
    from magic_pdf.data.dataset import PymuDocDataset
    MINERU_AVAILABLE = True
except ImportError as e:
    MINERU_AVAILABLE = False
    logger.warning(f"MinerU not available: {e}")

try:
    import fitz  # type: ignore[attr-defined]
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF (fitz) not available; PDF text extraction may be limited")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available")

try:
    import subprocess
    # Check for system-level antiword (prefer /usr/local/bin over pip package)
    antiword_check = subprocess.run(['which', '/usr/local/bin/antiword'], 
                                    capture_output=True)
    if antiword_check.returncode == 0:
        ANTIWORD_PATH = '/usr/local/bin/antiword'
        ANTIWORD_AVAILABLE = True
    else:
        # Fallback to PATH search
        antiword_check = subprocess.run(['which', 'antiword'], 
                                       capture_output=True)
        ANTIWORD_PATH = 'antiword'
        ANTIWORD_AVAILABLE = antiword_check.returncode == 0
except Exception:
    ANTIWORD_AVAILABLE = False
    ANTIWORD_PATH = 'antiword'

try:
    import textract  # type: ignore
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False
    logger.debug("textract not available")

try:
    from tika import parser as tika_parser
    TIKA_AVAILABLE = True
except ImportError:
    TIKA_AVAILABLE = False
    logger.debug("Apache Tika not available")

try:
    from striprtf.striprtf import rtf_to_text
    STRIPRTF_AVAILABLE = True
except ImportError:
    STRIPRTF_AVAILABLE = False
    logger.debug("striprtf not available - RTF file support limited")


class DocumentParser:
    """Parse various document formats to extract text content"""
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp(prefix="youtu_graphrag_")
        logger.info(f"DocumentParser initialized with temp dir: {self.temp_dir}")
        self.supports_multimodal = MINERU_AVAILABLE or PYMUPDF_AVAILABLE
    
    def parse_file(self, file_path: str, file_type: str) -> Optional[str]:
        """Parse a document file and extract plain text content."""
        text, _ = self._parse_with_multimodal(file_path, file_type)
        return text
    
    def parse_file_with_multimodal(self, file_path: str, file_type: str) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        """
        Parse a file and return both plain text and structured multimodal payload.
        """
        return self._parse_with_multimodal(file_path, file_type)

    def _parse_with_multimodal(self, file_path: str, file_type: str) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        file_type = file_type.lower()
        
        try:
            if file_type == '.pdf':
                return self._parse_pdf(file_path)
            elif file_type in ['.docx', '.doc']:
                return self._parse_docx(file_path)
            elif file_type in ['.txt', '.md']:
                text = self._parse_plain_text(file_path)
                payload = self._build_text_payload(file_path, text) if text else None
                return text, payload
            else:
                logger.warning(f"Unsupported file type for multimodal parsing: {file_type}")
                return None, None
        except Exception as e:
            logger.error(f"Error parsing {file_type} file: {e}")
            return None, None
    
    def _parse_pdf(self, pdf_path: str) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        """
        Parse PDF using MinerU (if available) with PyMuPDF fallbacks.
        Returns both text and structured payload.
        """
        payload = self._init_payload(pdf_path)
        text_parts: List[str] = []
        
        try:
            with open(pdf_path, 'rb') as f:
                pdf_bytes = f.read()
        except Exception as e:
            logger.error(f"Unable to read PDF file {pdf_path}: {e}")
            return None, None

        if MINERU_AVAILABLE:
            try:
                dataset = PymuDocDataset(pdf_bytes, lang='auto')
                for page_index in range(len(dataset)):
                    page_doc = dataset.get_page(page_index).get_doc()
                    page_payload = self._ensure_page_payload(payload, page_index + 1)
                    page_text = self._extract_page_blocks(page_doc, page_payload)
                    if page_text:
                        text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"MinerU structured extraction failed: {e}")
                payload = self._init_payload(pdf_path)
                text_parts = []

        if not text_parts and PYMUPDF_AVAILABLE:
            try:
                with fitz.open(pdf_path) as doc:  # type: ignore[attr-defined]
                    for page in doc:
                        page_payload = self._ensure_page_payload(payload, page.number + 1)
                        page_text = self._extract_page_blocks(page, page_payload)
                        if page_text:
                            text_parts.append(page_text)
            except Exception as e:
                logger.error(f"PyMuPDF extraction failed: {e}")

        if not text_parts:
            try:
                from pypdf import PdfReader  # type: ignore
                reader = PdfReader(pdf_path)
                for page_index, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                    except Exception as page_err:
                        logger.warning(
                            f"pypdf failed to extract page {page_index} of {pdf_path}: {page_err}"
                        )
                        continue
                    if page_text:
                        text_parts.append(page_text.strip())
                if text_parts and not payload.get("pages"):
                    payload = self._build_text_payload(pdf_path, "\n\n".join(text_parts))
            except ImportError:
                logger.debug("pypdf not installed; skipping pypdf fallback")
            except Exception as e:
                logger.error(f"pypdf fallback failed: {e}")

        if not text_parts:
            logger.error(f"Unable to extract text from PDF: {pdf_path}")
            return None, None
        
        extracted_text = '\n\n'.join(text_parts)
        logger.info(f"Successfully extracted {len(extracted_text)} chars from PDF")
        return extracted_text, self._finalize_payload(payload)
    
    def _parse_docx(self, docx_path: str) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        """
        Parse DOCX/DOC using available methods
        
        Args:
            docx_path: Path to DOCX/DOC file
            
        Returns:
            Extracted text content
        """
        file_ext = os.path.splitext(docx_path)[1].lower()
        
        # Check if the file is actually RTF (some .doc files are RTF in disguise)
        if self._is_rtf_file(docx_path):
            logger.info(f"Detected RTF file disguised as {file_ext}: {docx_path}")
            result = self._parse_rtf(docx_path)
            if result:
                return result, self._build_text_payload(docx_path, result)
        
        # Try python-docx first for .docx files
        if file_ext == '.docx' and DOCX_AVAILABLE:
            result = self._parse_with_python_docx(docx_path)
            if result:
                return result, self._build_text_payload(docx_path, result)
        
        # For .doc files or if python-docx fails, try alternative methods
        if file_ext == '.doc':
            # Priority 1: Try antiword first (fast, stable, no Python dependencies)
            if ANTIWORD_AVAILABLE:
                result = self._parse_with_antiword(docx_path)
                if result:
                    logger.info(f"Successfully extracted {len(result)} chars from DOC via antiword")
                    return result, self._build_text_payload(docx_path, result)
            
            # Priority 2: Try Apache Tika (best for WPS and complex formats)
            if TIKA_AVAILABLE:
                result = self._parse_with_tika(docx_path)
                if result:
                    logger.info(f"Successfully extracted {len(result)} chars from DOC via Apache Tika")
                    return result, self._build_text_payload(docx_path, result)
            
            # Priority 3: Try textract (if available, but has pip 24.1+ conflicts)
            if TEXTRACT_AVAILABLE:
                result = self._parse_with_textract(docx_path)
                if result:
                    logger.info(f"Successfully extracted {len(result)} chars from DOC via textract")
                    return result, self._build_text_payload(docx_path, result)
            
            # Priority 4: Try LibreOffice conversion (best for WPS/legacy formats)
            logger.debug(f"Trying LibreOffice for .doc file: {docx_path}")
            result = self._parse_doc_with_libreoffice(docx_path)
            if result:
                logger.info(f"Successfully extracted {len(result)} chars from DOC via LibreOffice")
                return result, self._build_text_payload(docx_path, result)
            else:
                logger.debug(f"LibreOffice parsing returned None for: {docx_path}")
        
        # Final fallback: try python-docx anyway (might work for some .doc files)
        if DOCX_AVAILABLE:
            result = self._parse_with_python_docx(docx_path)
            if result:
                return result, self._build_text_payload(docx_path, result)
        
        # Check file type to provide better error message
        file_type_hint = ""
        is_corrupted = False
        try:
            import subprocess
            file_info = subprocess.run(['file', docx_path], capture_output=True, text=True)
            if file_info.returncode == 0:
                info_lower = file_info.stdout.lower()
                if 'wps' in info_lower:
                    file_type_hint = " (WPS Office document)"
                    is_corrupted = True  # WPS 文档可能无法被 LibreOffice 解析
                elif 'composite document' in info_lower or 'ole' in info_lower:
                    if 'microsoft' in info_lower:
                        file_type_hint = " (Legacy Microsoft Word format)"
                    else:
                        file_type_hint = " (OLE document)"
        except Exception as e:
            logger.warning(f"Failed to determine file type for {docx_path}: {e}")
        
        logger.error(f"Unable to parse {file_ext} file{file_type_hint}: {docx_path}")
        
        if is_corrupted:
            logger.warning("⚠️  This WPS Office document cannot be parsed by available tools")
            logger.info("Recommended solutions:")
            logger.info("  1. Open in WPS Office and save as .docx format")
            logger.info("  2. Open in Microsoft Word and save as .docx format")
            logger.info("  3. Use online converters (e.g., zamzar.com, cloudconvert.com)")
        else:
            logger.info("Recommended solutions:")
            logger.info("  1. Install LibreOffice (best compatibility): sudo yum install libreoffice-headless")
            logger.info("  2. Convert file to .docx format using Microsoft Word or WPS")
            logger.info("  3. For standard .doc files: Install antiword (sudo apt-get install antiword)")
        return None, None
    
    def _parse_with_python_docx(self, docx_path: str) -> Optional[str]:
        """Parse using python-docx library"""
        if not DOCX_AVAILABLE:
            return None
        
        try:
            doc = DocxDocument(docx_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            extracted_text = '\n'.join(text_parts)
            
            if not extracted_text.strip():
                return None
            
            logger.info(f"Successfully extracted {len(extracted_text)} chars via python-docx")
            return extracted_text
            
        except Exception as e:
            logger.debug(f"python-docx failed: {e}")
            return None
    
    def _parse_with_textract(self, doc_path: str) -> Optional[str]:
        """Parse using textract library"""
        if not TEXTRACT_AVAILABLE:
            return None
        
        try:
            text = textract.process(doc_path).decode('utf-8')
            if text and text.strip():
                return text.strip()
        except Exception as e:
            logger.debug(f"textract failed: {e}")
        return None

    # -------------------- Multimodal helpers -------------------- #
    def _parse_plain_text(self, file_path: str) -> Optional[str]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Failed to read plain text file {file_path}: {e}")
            return None

    def _init_payload(self, source_path: str) -> Dict[str, Any]:
        return {
            "source_file": os.path.basename(source_path),
            "pages": []
        }

    def _ensure_page_payload(self, payload: Dict[str, Any], page_number: int) -> Dict[str, Any]:
        for page in payload["pages"]:
            if page["page_number"] == page_number:
                return page
        page_payload = {
            "page_number": page_number,
            "text_spans": [],
            "tables": [],
            "formulas": [],
            "images": []
        }
        payload["pages"].append(page_payload)
        return page_payload

    def _finalize_payload(self, payload: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        if not payload["pages"]:
            return None
        for page in payload["pages"]:
            if any(page[key] for key in ("text_spans", "tables", "formulas", "images")):
                return payload
        return None

    def _extract_page_blocks(self, page_doc, page_payload: Dict[str, Any]) -> Optional[str]:
        page_text_parts: List[str] = []
        try:
            blocks = page_doc.get_text("blocks")
        except Exception as e:
            logger.debug(f"Failed to read page blocks: {e}")
            blocks = []

        for block in blocks:
            if not isinstance(block, (list, tuple)) or len(block) < 5:
                continue
            x0, y0, x1, y1, text = block[:5]
            cleaned = str(text).strip()
            if not cleaned:
                continue
            entry = {"text": cleaned, "bbox": [float(x0), float(y0), float(x1), float(y1)]}
            if self._looks_like_table(cleaned):
                page_payload["tables"].append(entry)
            elif self._looks_like_formula(cleaned):
                page_payload["formulas"].append(entry)
            else:
                page_payload["text_spans"].append(entry)
            page_text_parts.append(cleaned)

        page_payload["images"].extend(self._extract_images_from_page(page_doc))
        return "\n".join(page_text_parts) if page_text_parts else None

    def _extract_images_from_page(self, page_doc) -> List[Dict[str, Any]]:
        images: List[Dict[str, Any]] = []
        if not PYMUPDF_AVAILABLE:
            return images
        try:
            for img in page_doc.get_images(full=True):
                xref = img[0]
                bbox = None
                rects = []
                for rect in page_doc.get_image_rects(xref):
                    rects.append(self._normalize_bbox(rect))
                if rects:
                    bbox = rects[0]
                img_bytes = None
                ext = None
                try:
                    base_image = page_doc.parent.extract_image(xref)
                    img_bytes = base_image.get("image")
                    ext = base_image.get("ext")
                except Exception as e:
                    logger.debug(f"Failed to extract image bytes for xref {xref}: {e}")
                images.append({
                    "bbox": bbox,
                    "bytes": img_bytes,
                    "ext": ext,
                    "caption_candidates": self._caption_from_rect(page_doc, bbox)
                })
        except Exception as e:
            logger.debug(f"Failed to extract images on page: {e}")
        return images

    def _normalize_bbox(self, rect) -> Optional[List[float]]:
        if rect is None:
            return None
        if hasattr(rect, "x0"):
            return [float(rect.x0), float(rect.y0), float(rect.x1), float(rect.y1)]
        if isinstance(rect, (list, tuple)) and len(rect) >= 4:
            return [float(rect[0]), float(rect[1]), float(rect[2]), float(rect[3])]
        return None

    def _caption_from_rect(self, page_doc, bbox: Optional[List[float]]) -> List[str]:
        if not bbox:
            return []
        try:
            blocks = page_doc.get_text("blocks")
        except Exception:
            return []
        captions: List[str] = []
        x0, y0, x1, y1 = bbox
        for block in blocks:
            if not isinstance(block, (list, tuple)) or len(block) < 5:
                continue
            bx0, by0, bx1, by1, text = block[:5]
            snippet = str(text).strip()
            if not snippet:
                continue
            horizontal_overlap = not (bx1 < x0 or bx0 > x1)
            vertical_distance = min(abs(by0 - y1), abs(y0 - by1))
            if horizontal_overlap and vertical_distance < 60:
                captions.append(snippet)
            if len(captions) >= 2:
                break
        return captions

    def _looks_like_table(self, text: str) -> bool:
        lines = [line for line in text.splitlines() if line.strip()]
        if len(lines) < 2:
            return False
        delimiter_hits = sum(1 for line in lines if any(delim in line for delim in TABLE_DELIMITERS))
        numeric_lines = sum(1 for line in lines if sum(ch.isdigit() for ch in line) >= 4)
        if delimiter_hits >= max(1, len(lines) // 2):
            return True
        return numeric_lines >= 2 and len(lines[0].split()) >= 3

    def _looks_like_formula(self, text: str) -> bool:
        if len(text) > 400:
            return False
        if any(ch in FORMULA_HINT_CHARS for ch in text):
            return True
        if re.search(r"[A-Za-z]\\s*=\\s*[0-9A-Za-z]", text):
            return True
        if "\\" in text and ("frac" in text or "sum" in text):
            return True
        return False

    def _build_text_payload(self, source_path: str, text: str) -> Optional[Dict[str, Any]]:
        payload = self._init_payload(source_path)
        page_payload = self._ensure_page_payload(payload, 1)
        for paragraph in re.split(r"\n{2,}", text):
            cleaned = paragraph.strip()
            if not cleaned:
                continue
            entry = {"text": cleaned, "bbox": None}
            if self._looks_like_table(cleaned):
                page_payload["tables"].append(entry)
            elif self._looks_like_formula(cleaned):
                page_payload["formulas"].append(entry)
            else:
                page_payload["text_spans"].append(entry)
        return self._finalize_payload(payload)
    
    def _parse_with_tika(self, doc_path: str) -> Optional[str]:
        """
        Parse using Apache Tika (supports WPS, legacy Word, and many other formats)
        
        Args:
            doc_path: Path to the document file
            
        Returns:
            Extracted text content
        """
        if not TIKA_AVAILABLE:
            return None
        
        try:
            # Apache Tika 可以处理几乎所有文档格式,包括 WPS Office 文档
            parsed = tika_parser.from_file(doc_path)
            text = parsed.get('content', '')
            
            if text and text.strip():
                return text.strip()
            else:
                logger.debug(f"Tika returned empty content for: {doc_path}")
                return None
                
        except Exception as e:
            logger.debug(f"Apache Tika parsing failed: {e}")
            return None
    
    def _parse_with_antiword(self, doc_path: str) -> Optional[str]:
        """Parse using antiword command-line tool"""
        if not ANTIWORD_AVAILABLE:
            return None
        
        try:
            result = subprocess.run(
                [ANTIWORD_PATH, doc_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except Exception as e:
            logger.debug(f"antiword failed: {e}")
        return None
    
    def _parse_doc_with_libreoffice(self, doc_path: str) -> Optional[str]:
        """Convert .doc to .txt using LibreOffice and read the result"""
        try:
            # Check if libreoffice is available
            lo_check = subprocess.run(
                ['which', 'libreoffice'],
                capture_output=True
            )
            if lo_check.returncode != 0:
                logger.debug("LibreOffice not found in PATH")
                return None
            
            logger.debug(f"LibreOffice found, attempting conversion for: {doc_path}")
            
            # Create temp directory for conversion
            import shutil
            temp_dir = tempfile.mkdtemp(prefix="doc_convert_")
            
            try:
                # Copy file to temp dir with safe filename (避免特殊字符导致转换失败)
                import hashlib
                file_ext = os.path.splitext(doc_path)[1]
                # 使用文件路径的哈希生成安全的文件名
                safe_name = hashlib.md5(doc_path.encode()).hexdigest()
                temp_doc = os.path.join(temp_dir, f"{safe_name}{file_ext}")
                shutil.copy2(doc_path, temp_doc)
                
                logger.debug(f"Copied to temp file: {temp_doc}")
                
                # Convert to txt
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'txt:Text', 
                     '--outdir', temp_dir, temp_doc],
                    capture_output=True,
                    text=True,
                    timeout=60,  # 增加超时时间,处理大文件
                    check=False  # 不抛出异常,稍后检查返回码
                )
                
                logger.debug(f"LibreOffice exit code: {result.returncode}")
                if result.stdout:
                    logger.debug(f"LibreOffice stdout: {result.stdout[:200]}")
                if result.stderr:
                    logger.debug(f"LibreOffice stderr: {result.stderr[:200]}")
                
                # Check if conversion succeeded
                if result.returncode != 0:
                    logger.debug(f"LibreOffice conversion failed with code {result.returncode}")
                    return None
                
                # Read the converted file
                txt_path = os.path.join(temp_dir, f"{safe_name}.txt")
                logger.debug(f"Looking for output file: {txt_path}")
                
                if os.path.exists(txt_path):
                    with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                    logger.debug(f"Read {len(text)} chars from converted file")
                    if text.strip():
                        return text.strip()
                    else:
                        logger.debug("Converted file is empty")
                        return None
                else:
                    # 即使返回码为0,如果输出文件不存在,说明转换实际失败了
                    logger.debug(f"LibreOffice output file not found: {txt_path}")
                    # 检查stderr中是否有"no export filter"或其他错误
                    if "no export filter" in result.stderr or "Error:" in result.stderr:
                        logger.warning(f"LibreOffice cannot parse this document format. Error: {result.stderr[:150]}")
                    return None
            finally:
                # Cleanup
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
                    
        except Exception as e:
            logger.debug(f"LibreOffice conversion exception: {e}")
        return None
    
    def _is_rtf_file(self, file_path: str) -> bool:
        """
        Check if a file is actually RTF format (regardless of extension)
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            True if file is RTF format
        """
        try:
            with open(file_path, 'rb') as f:
                # RTF files start with {\rtf
                header = f.read(10)
                return header.startswith(b'{\\rtf')
        except Exception as e:
            logger.debug(f"Error checking RTF format: {e}")
            return False
    
    def _parse_rtf(self, rtf_path: str) -> Optional[str]:
        """
        Parse RTF file and extract text content
        
        Args:
            rtf_path: Path to RTF file
            
        Returns:
            Extracted text content
        """
        # Method 1: Try striprtf library if available
        if STRIPRTF_AVAILABLE:
            try:
                with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
                    rtf_content = f.read()
                text = rtf_to_text(rtf_content)
                if text and text.strip():
                    logger.info(f"Successfully extracted {len(text)} chars from RTF via striprtf")
                    return text.strip()
            except Exception as e:
                logger.debug(f"striprtf parsing failed: {e}")
        
        # Method 2: Try textract if available
        if TEXTRACT_AVAILABLE:
            try:
                text = textract.process(rtf_path).decode('utf-8', errors='ignore')
                if text and text.strip():
                    logger.info(f"Successfully extracted {len(text)} chars from RTF via textract")
                    return text.strip()
            except Exception as e:
                logger.debug(f"textract RTF parsing failed: {e}")
        
        # Method 3: Try LibreOffice conversion
        try:
            # Check if libreoffice is available
            lo_check = subprocess.run(
                ['which', 'libreoffice'],
                capture_output=True
            )
            if lo_check.returncode == 0:
                import hashlib
                import shutil
                temp_dir = tempfile.mkdtemp(prefix="rtf_convert_")
                try:
                    # Copy file to temp dir with safe filename
                    file_ext = os.path.splitext(rtf_path)[1]
                    safe_name = hashlib.md5(rtf_path.encode()).hexdigest()
                    temp_rtf = os.path.join(temp_dir, f"{safe_name}{file_ext}")
                    shutil.copy2(rtf_path, temp_rtf)
                    
                    result = subprocess.run(
                        ['libreoffice', '--headless', '--convert-to', 'txt:Text',
                         '--outdir', temp_dir, temp_rtf],
                        capture_output=True,
                        text=True,
                        timeout=60,
                        check=False
                    )
                    
                    if result.returncode == 0:
                        txt_path = os.path.join(temp_dir, f"{safe_name}.txt")
                        
                        if os.path.exists(txt_path):
                            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                                text = f.read()
                            if text.strip():
                                logger.info(f"Successfully extracted {len(text)} chars from RTF via LibreOffice")
                                return text.strip()
                finally:
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
        except Exception as e:
            logger.debug(f"LibreOffice RTF conversion failed: {e}")
        
        logger.error(f"Unable to parse RTF file: {rtf_path}")
        logger.info("Hint: Install 'striprtf' for RTF support: pip install striprtf")
        logger.info("  Or install LibreOffice: sudo yum install libreoffice-headless")
        return None
    
    def cleanup(self):
        """Clean up temporary files"""
        try:
            import shutil
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                logger.info(f"Cleaned up temp dir: {self.temp_dir}")
        except Exception as e:
            logger.error(f"Error cleaning up temp dir: {e}")


# Global parser instance
_parser_instance = None

def get_parser() -> DocumentParser:
    """Get or create global parser instance"""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = DocumentParser()
    return _parser_instance
